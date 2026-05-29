import os
import json
import datetime
import openpyxl
from utils.logger import get_logger
from utils.dedup import get_next_uid

import config

logger = get_logger('db_manager')

HEADERS = [
    'UID', 'Date', 'Source', 'Expression', 'POS',
    'Pronunciation', 'Meaning_KR', 'Original_Text', 'Applied_Example'
]

def _generate_uid(date_str: str, sequence_num: int) -> str:
    """Generate a unique identifier for an expression entry: CHI-YYYYMMDD-NNN"""
    date_compact = date_str.replace('-', '')
    return f"CHI-{date_compact}-{sequence_num:03d}"

def _prepare_rows(expressions: list[dict], date_str: str, index_data: dict) -> list[list]:
    rows = []
    for expr in expressions:
        seq_num = get_next_uid(index_data, date_str)
        uid = _generate_uid(date_str, seq_num)
        
        # Attach UID back to the dict
        expr['uid'] = uid

        row = [
            uid,
            date_str,
            expr.get('source', 'Inspiration'),
            expr.get('expression', ''),
            expr.get('pos', ''),
            expr.get('pinyin', ''),
            expr.get('meaning_kr', ''),
            expr.get('original_text', ''),
            expr.get('applied_example', '')
        ]
        rows.append(row)
    return rows

def _save_master_excel(rows: list[list], headers: list[str]) -> int:
    file_path = config.EXCEL_FILENAME
    os.makedirs(config.DB_DIR, exist_ok=True)
    try:
        if os.path.exists(file_path):
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            logger.info(f"Opened master Excel: {file_path} (last row: {ws.max_row})")
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Expressions"
            ws.append(headers)
            logger.info(f"Created new master Excel: {file_path}")

        for row in rows:
            ws.append(row)

        wb.save(file_path)
        logger.info(f"Appended {len(rows)} expressions to master Excel.")
        return len(rows)
    except Exception as e:
        logger.error(f"Failed to save master Excel: {e}")
        raise

def _save_daily_excel(rows: list[list], headers: list[str], date_str: str) -> None:
    os.makedirs(config.DAILY_DIR, exist_ok=True)
    file_path = os.path.join(config.DAILY_DIR, f"Expressions_{date_str}.xlsx")
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"Daily_{date_str}"
        ws.append(headers)

        for row in rows:
            ws.append(row)

        wb.save(file_path)
        logger.info(f"Saved daily Excel sheet: {file_path}")
    except Exception as e:
        logger.error(f"Failed to save daily Excel: {e}")
        raise

def _save_word_note(expressions: list[dict], date_str: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    os.makedirs(config.PRINT_DIR, exist_ok=True)
    file_path = os.path.join(config.PRINT_DIR, f"Study_Note_{date_str}.docx")

    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

        def _set_font(run, name="맑은 고딕", ch_name="Microsoft YaHei", size=11, bold=False, italic=False, color=None):
            run.font.name = name
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)
            r = run._r
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), ch_name)
            rFonts.set(qn('w:ascii'), name)
            rPr.append(rFonts)

        # Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after  = Pt(6)
        title_r = title_p.add_run(f"📝 Daily Chinese Study Note  |  {date_str}")
        _set_font(title_r, size=15, bold=True, color=(180, 40, 40))

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_before = Pt(0)
        sub_p.paragraph_format.space_after  = Pt(14)
        sub_r = sub_p.add_run(f"오늘의 핵심 네이티브 표현  {len(expressions)}개  ·  한어 병음 · 한국어 의미 · 원문 · 실전 예문 수록")
        _set_font(sub_r, size=9, italic=True, color=(100, 100, 100))

        # Cards
        for i, expr in enumerate(expressions, 1):
            expression = expr.get('expression', '')
            pos        = expr.get('pos', '').upper()
            pinyin     = expr.get('pinyin', '')
            meaning    = expr.get('meaning_kr', '')
            orig_txt   = expr.get('original_text', '')
            appl_ex    = expr.get('applied_example', '')

            # Expression Line
            head_p = doc.add_paragraph()
            head_p.paragraph_format.space_before = Pt(10)
            head_p.paragraph_format.space_after  = Pt(2)

            num_r = head_p.add_run(f"{i:03d}.  ")
            _set_font(num_r, size=10, color=(150, 150, 150))

            expr_r = head_p.add_run(expression)
            _set_font(expr_r, name="Microsoft YaHei", size=14, bold=True, color=(150, 20, 20))

            pos_r = head_p.add_run(f"   [{pos}]")
            _set_font(pos_r, name="Calibri", size=9, color=(120, 120, 120))

            # Pinyin
            pinyin_p = doc.add_paragraph()
            pinyin_p.paragraph_format.space_before = Pt(0)
            pinyin_p.paragraph_format.space_after  = Pt(2)
            pin_label = pinyin_p.add_run("🔊 ")
            _set_font(pin_label, size=10)
            pin_r = pinyin_p.add_run(pinyin)
            _set_font(pin_r, name="Calibri", size=10, italic=True, color=(80, 80, 80))

            # Meaning
            mean_p = doc.add_paragraph()
            mean_p.paragraph_format.space_before = Pt(0)
            mean_p.paragraph_format.space_after  = Pt(3)
            mean_label = mean_p.add_run("💡 의미  ")
            _set_font(mean_label, size=10, bold=True, color=(40, 40, 40))
            mean_r = mean_p.add_run(meaning)
            _set_font(mean_r, size=10.5)

            # Original
            orig_p = doc.add_paragraph()
            orig_p.paragraph_format.space_before = Pt(0)
            orig_p.paragraph_format.space_after  = Pt(2)
            orig_label = orig_p.add_run("📌 원문  ")
            _set_font(orig_label, size=9.5, bold=True, color=(100, 100, 100))
            orig_r = orig_p.add_run(f'"{orig_txt}"')
            _set_font(orig_r, name="Microsoft YaHei", size=9.5, italic=True, color=(90, 90, 90))

            # Example
            ex_p = doc.add_paragraph()
            ex_p.paragraph_format.space_before = Pt(2)
            ex_p.paragraph_format.space_after  = Pt(2)
            ex_label = ex_p.add_run("✏️ 예문  ")
            _set_font(ex_label, size=10, bold=True, color=(180, 40, 40))
            ex_r = ex_p.add_run(appl_ex)
            _set_font(ex_r, name="Microsoft YaHei", size=10.5, bold=True, color=(140, 30, 30))

            if i < len(expressions):
                sep_p = doc.add_paragraph()
                sep_p.paragraph_format.space_before = Pt(6)
                sep_p.paragraph_format.space_after  = Pt(0)
                sep_r = sep_p.add_run("─" * 60)
                _set_font(sep_r, size=7, color=(200, 200, 200))

        doc.save(file_path)
        logger.info(f"Saved daily study Word note: {file_path}")
    except Exception as e:
        logger.error(f"Failed to save Word note: {e}")
        raise

def _save_markdown_note(expressions: list[dict], date_str: str) -> None:
    os.makedirs(config.PRINT_DIR, exist_ok=True)
    file_path = os.path.join(config.PRINT_DIR, f"Study_Note_{date_str}.md")

    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"# 📝 Daily Chinese Native Expressions Study Note ({date_str})\n\n")
            f.write(f"오늘 학습할 핵심 중국어 표현 **{len(expressions)}개**입니다.\n")
            f.write("눈으로 소리 내어 읽고, 아래 한어 병음과 실전 예문을 보며 내 손으로 직접 뜻과 쓰임새를 익혀보세요!\n\n")
            
            f.write("## 📂 학습 표현 목록\n")
            f.write("| 번호 | 표현 (Expression) | 품사 (POS) | 의미 (Meaning) | 병음 (Pinyin) |\n")
            f.write("| :---: | :--- | :--- | :--- | :--- |\n")
            for i, expr in enumerate(expressions, 1):
                f.write(f"| {i:03d} | **{expr.get('expression', '')}** | `{expr.get('pos', '').lower()}` | {expr.get('meaning_kr', '')} | {expr.get('pinyin', '')} |\n")
            f.write("\n---\n\n")

            f.write("## 🔍 세부 표현 및 실전 예문 학습\n\n")
            for i, expr in enumerate(expressions, 1):
                f.write(f"### {i}. **{expr.get('expression', '')}** `[{expr.get('pos', '').lower()}]`\n")
                f.write(f"*   **병음**: {expr.get('pinyin', '')}\n")
                f.write(f"*   **의미**: {expr.get('meaning_kr', '')}\n")
                f.write(f"*   **원문 Context**: *\"{expr.get('original_text', '')}\"*\n")
                f.write(f"*   **실전 예문 (Applied Example)**:\n")
                f.write(f"    > **{expr.get('applied_example', '')}**\n\n")
                f.write("---\n\n")
        logger.info(f"Saved daily study Markdown note: {file_path}")
    except Exception as e:
        logger.error(f"Failed to save Markdown note: {e}")
        raise

def save_expressions(expressions: list[dict], index_data: dict) -> int:
    if not expressions:
        logger.warning("No expressions to save")
        return 0
    date_str = datetime.date.today().strftime('%Y-%m-%d')

    rows = _prepare_rows(expressions, date_str, index_data)
    logger.info(f"Prepared {len(rows)} rows for saving (date: {date_str})")

    # 1. Excel Master DB
    saved_count = _save_master_excel(rows, HEADERS)
    
    # 2. Daily Excel
    _save_daily_excel(rows, HEADERS, date_str)
    
    # 3. Daily DOCX Note
    _save_word_note(expressions, date_str)
    
    # 4. Daily MD Note
    _save_markdown_note(expressions, date_str)

    return saved_count
